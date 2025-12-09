#!/usr/bin/env python3
"""将 docx 文件转换为 markdown 格式，保留表格，去掉双引号和**"""
import re
from datetime import datetime
from docx import Document
from pathlib import Path

def docx_to_markdown(docx_path):
    """读取 docx 文件并转换为 markdown"""
    doc = Document(docx_path)
    markdown_lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_lines.append("")
            continue
        
        # 检查是否是标题
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
            markdown_lines.append(f"{'#' * level} {text}")
        else:
            # 普通段落
            markdown_lines.append(text)
    
    # 处理表格
    for table in doc.tables:
        markdown_lines.append("")  # 空行分隔
        # 表头
        header_row = table.rows[0]
        header_cells = [cell.text.strip() for cell in header_row.cells]
        markdown_lines.append("| " + " | ".join(header_cells) + " |")
        # 分隔行
        markdown_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        # 数据行
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            markdown_lines.append("| " + " | ".join(cells) + " |")
        markdown_lines.append("")  # 空行分隔
    
    # 合并所有行
    markdown_text = "\n".join(markdown_lines)
    
    # 去掉所有类型的引号
    markdown_text = markdown_text.replace('"', '')
    markdown_text = markdown_text.replace('"', '')
    markdown_text = markdown_text.replace('"', '')
    markdown_text = markdown_text.replace('"', '')
    markdown_text = markdown_text.replace('"', '')
    markdown_text = markdown_text.replace('"', '')
    markdown_text = markdown_text.replace('"', '')
    markdown_text = markdown_text.replace('"', '')
    
    # 去掉 **（加粗标记）
    markdown_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', markdown_text)
    
    return markdown_text

def create_markdown_file(title, content, category, tags, output_dir):
    """创建完整的 markdown 文件，包含 front matter"""
    # 生成文件名（基于标题）
    filename = title.replace(" ", "-").replace("：", "-").replace(":", "-")
    filename = re.sub(r'[^\w\-]', '', filename)
    filename = filename.lower()[:50]  # 限制长度
    filename = f"{category}-{filename}.markdown"
    
    # 获取当前日期
    date_str = datetime.now().strftime("%Y-%m-%d")
    
    # 构建 front matter
    front_matter = f"""---
title: {title}
date: {date_str}
tags: {tags}
category: {category}
---

"""
    
    # 组合完整内容
    full_content = front_matter + content
    
    # 保存文件
    output_path = output_dir / filename
    output_path.write_text(full_content, encoding="utf-8")
    print(f"已保存: {output_path}")
    return output_path

if __name__ == "__main__":
    source_dir = Path("/Users/tuboshu/Desktop/我的网站")
    output_dir = Path("/Users/tuboshu/土拨鼠的蛋/inflammaging_site/markdown 文章")
    
    # 找到两个文件
    qiaoyong_file = None
    zoster_file = None
    
    for file in source_dir.glob("*.docx"):
        if "巧用" in file.name and not file.name.startswith("."):
            qiaoyong_file = file
        elif "带状疱疹" in file.name and not file.name.startswith("."):
            zoster_file = file
    
    # 处理第一个文件：巧用 -> papers
    if qiaoyong_file:
        print(f"\n处理文件: {qiaoyong_file.name}")
        content = docx_to_markdown(qiaoyong_file)
        # 从内容中提取标题（第一个一级标题）
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        title = title_match.group(1) if title_match else "巧用自然实验：解析回归不连续设计研究方法"
        # 去掉标题行（因为会在 front matter 中）
        content = re.sub(r'^#\s+.+$', '', content, count=1, flags=re.MULTILINE).strip()
        # 重新添加标题（作为正文标题）
        content = f"# {title}\n\n{content}"
        create_markdown_file(
            title=title,
            content=content,
            category="papers",
            tags="[Research methods, Regression discontinuity design, Natural experiment, Causal inference]",
            output_dir=output_dir
        )
    
    # 处理第二个文件：带状疱疹疫苗 -> pathways
    if zoster_file:
        print(f"\n处理文件: {zoster_file.name}")
        content = docx_to_markdown(zoster_file)
        # 从内容中提取标题（第一个一级标题）
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        title = title_match.group(1) if title_match else "带状疱疹疫苗接种对不同阶段痴呆症影响的研究摘要"
        # 去掉标题行（因为会在 front matter 中）
        content = re.sub(r'^#\s+.+$', '', content, count=1, flags=re.MULTILINE).strip()
        # 重新添加标题（作为正文标题）
        content = f"# {title}\n\n{content}"
        create_markdown_file(
            title=title,
            content=content,
            category="pathways",
            tags="[Herpes zoster vaccine, Alzheimer's disease, Dementia, Neuroinflammation, Prevention]",
            output_dir=output_dir
        )
    
    print("\n转换完成！")
